from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

BASE = Path(__file__).resolve().parents[1]
OUT_DIR = BASE / 'docs'
OUT_DIR.mkdir(parents=True, exist_ok=True)

PROJECT_NAME = 'RiskGuard'
SUBTITLE = 'Aplicacion Web para Evaluacion y Tratamiento de Riesgos de Seguridad de la Informacion'

TEAM = [
    'Josue Ayala',
    'Emil Encalada',
    'Daniel Mena',
    'Jorge Ramos',
]

COURSE = 'Seguridad Informatica'
DOCENTE = '[Docente]'
UNIVERSIDAD = '[Universidad]'


def set_normal_style(doc: Document) -> None:
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    # Compact spacing
    pformat = style.paragraph_format
    pformat.space_after = Pt(6)
    pformat.space_before = Pt(0)
    pformat.line_spacing = 1.0


def add_title_page(doc: Document, title: str, subtitle: str) -> None:
    set_normal_style(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(UNIVERSIDAD)
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(COURSE)
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.italic = True
    run.font.size = Pt(12)

    doc.add_paragraph('')

    doc.add_paragraph('Integrantes:').runs[0].bold = True
    for member in TEAM:
        doc.add_paragraph(f'- {member}')

    doc.add_paragraph('')

    doc.add_paragraph(f'Docente: {DOCENTE}')
    doc.add_paragraph(f'Fecha: {date.today().isoformat()}')

    doc.add_page_break()


def manual_usuario() -> Path:
    doc = Document()
    add_title_page(doc, 'Manual de Usuario', f'{PROJECT_NAME} - {SUBTITLE}')

    doc.add_heading('1. Que es RiskGuard', level=1)
    doc.add_paragraph(
        'RiskGuard es una aplicacion web sencilla para registrar activos, amenazas, vulnerabilidades y '
        'evaluar riesgos usando una metodologia de probabilidad x impacto. '
        'Tambien permite proponer controles (con referencia ISO/IEC 27002:2022), registrar el tratamiento, '
        'calcular el riesgo residual, llevar un seguimiento con KPIs y generar un reporte en PDF.'
    )

    doc.add_heading('2. Requisitos', level=1)
    doc.add_paragraph('Para usar RiskGuard necesitas:')
    for item in [
        'Windows / macOS / Linux',
        'Python 3.10+ (recomendado 3.11)',
        'Navegador web (Chrome/Firefox/Edge)',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3. Instalacion rapida', level=1)
    steps = [
        'Descargar y descomprimir el archivo RiskGuard.zip.',
        'Abrir una terminal y entrar a la carpeta: RiskGuard/app',
        'Crear entorno virtual: python -m venv .venv',
        'Activar entorno virtual:',
        '  - Windows: .venv\\Scripts\\activate',
        '  - macOS/Linux: source .venv/bin/activate',
        'Instalar dependencias: pip install -r ../requirements.txt',
        'Ejecutar la app: python app.py',
        'Abrir en el navegador: http://127.0.0.1:5000',
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    doc.add_heading('4. Uso del sistema (flujo recomendado)', level=1)
    doc.add_paragraph('Para una demo completa en clase, sigue este orden:')

    doc.add_paragraph('4.1 Registrar Activos', style=None).runs[0].bold = True
    doc.add_paragraph('Menu: Activos -> + Nuevo activo', style='List Bullet')
    doc.add_paragraph('Completa nombre, tipo, proceso/area y la valoracion CID (1-3).', style='List Bullet')
    doc.add_paragraph('CAPTURA: Pantalla de registro de activos (Figura 1)', style='Intense Quote')

    doc.add_paragraph('4.2 Registrar Amenazas y Vulnerabilidades', style=None).runs[0].bold = True
    doc.add_paragraph('Menu: Amenazas / Vulnerabilidades -> + Nuevo', style='List Bullet')
    doc.add_paragraph('Recomendado: usar categorias (ej: Malware, Ingenieria social, Configuracion).', style='List Bullet')
    doc.add_paragraph('CAPTURA: Catalogos (Figura 2)', style='Intense Quote')

    doc.add_paragraph('4.3 Registrar Controles', style=None).runs[0].bold = True
    doc.add_paragraph('Menu: Controles -> + Nuevo control', style='List Bullet')
    doc.add_paragraph('Puedes anotar referencia ISO/IEC 27002:2022 (ej: "5.23" o "Control de accesos").', style='List Bullet')
    doc.add_paragraph('CAPTURA: Controles (Figura 3)', style='Intense Quote')

    doc.add_paragraph('4.4 Crear Riesgos (escenarios)', style=None).runs[0].bold = True
    doc.add_paragraph('Menu: Riesgos -> + Nuevo riesgo', style='List Bullet')
    doc.add_paragraph('Selecciona Activo + Amenaza + Vulnerabilidad, define Probabilidad (1-5).', style='List Bullet')
    doc.add_paragraph('El impacto se toma del activo (por CID), o puedes usar Impacto (override).', style='List Bullet')
    doc.add_paragraph('CAPTURA: Crear riesgo (Figura 4)', style='Intense Quote')

    doc.add_paragraph('4.5 Tratamiento y plan de accion', style=None).runs[0].bold = True
    doc.add_paragraph('En el detalle del riesgo: clic en "Tratamiento"', style='List Bullet')
    doc.add_paragraph('Elige estrategia (Mitigar/Transferir/Aceptar/Evitar), controles propuestos, responsable, fecha y estado.', style='List Bullet')
    doc.add_paragraph('Si eliges Aceptar, llena justificacion y aprobado por.', style='List Bullet')
    doc.add_paragraph('CAPTURA: Tratamiento (Figura 5)', style='Intense Quote')

    doc.add_paragraph('4.6 Riesgo residual', style=None).runs[0].bold = True
    doc.add_paragraph('En el detalle del riesgo: clic en "Residual"', style='List Bullet')
    doc.add_paragraph('Registra P/I residual despues de implementar controles.', style='List Bullet')
    doc.add_paragraph('CAPTURA: Residual (Figura 6)', style='Intense Quote')

    doc.add_paragraph('4.7 Reporte PDF', style=None).runs[0].bold = True
    doc.add_paragraph('Menu: Reportes -> se descarga un PDF con inventario, matriz de riesgos, tratamientos y KPIs.', style='List Bullet')

    doc.add_heading('5. Solucion de problemas (rapido)', level=1)
    doc.add_paragraph('Si no carga el sistema: revisa que la terminal diga "Running on http://127.0.0.1:5000".', style='List Bullet')
    doc.add_paragraph('Si falla una dependencia: ejecuta pip install -r ../requirements.txt.', style='List Bullet')
    doc.add_paragraph('Si quieres datos de ejemplo: ejecuta python seed.py (desde la carpeta app).', style='List Bullet')

    out = OUT_DIR / 'Manual_Usuario_RiskGuard.docx'
    doc.save(out)
    return out


def documento_tecnico() -> Path:
    doc = Document()
    add_title_page(doc, 'Documento Tecnico', f'{PROJECT_NAME} - {SUBTITLE}')

    doc.add_heading('1. Introduccion', level=1)
    doc.add_paragraph(
        'La seguridad de la informacion no es solo "poner antivirus". En una organizacion real, '
        'siempre existen activos (datos, sistemas, procesos) que pueden ser afectados por amenazas, '
        'explotando vulnerabilidades. Por eso se necesita una metodologia clara para identificar, '
        'valorar, tratar y monitorear riesgos. '        'Este documento describe el diseno e implementacion de RiskGuard, una aplicacion web ligera '
        'creada para apoyar el proceso de gestion de riesgos con una interfaz intuitiva, validacion de datos '
        'y un reporte automatico para comunicacion con partes interesadas.'
    )

    doc.add_heading('2. Objetivo y alcance', level=1)
    doc.add_paragraph('Objetivo general:')
    doc.add_paragraph(
        'Desarrollar una aplicacion web que permita evaluar riesgos de seguridad de la informacion usando '
        'una metodologia basada en Probabilidad x Impacto, incorporando tratamiento, riesgo residual, '
        'monitoreo con KPIs y generacion de reportes.',
        style='List Bullet'
    )

    doc.add_paragraph('Alcance:')
    for s in [
        'Inventario de activos con valoracion CID (Confidencialidad, Integridad, Disponibilidad).',
        'Catalogos de amenazas y vulnerabilidades.',
        'Registro de escenarios de riesgo (Activo + Amenaza + Vulnerabilidad).',
        'Tratamiento del riesgo (estrategia, controles, responsable, fecha limite y estado).',
        'Evaluacion del riesgo residual despues de implementar controles.',
        'Monitoreo mediante dashboard con KPIs y registro de incidentes.',
        'Generacion de reporte PDF para comunicacion con interesados.',
    ]:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('3. Metodologia de evaluacion', level=1)
    doc.add_heading('3.1 Valoracion del activo (CID)', level=2)
    doc.add_paragraph(
        'Cada activo se califica con tres criterios: Confidencialidad (C), Integridad (I) y Disponibilidad (D) '
        'en una escala de 1 a 3. El sistema calcula CID total = C + I + D (3 a 9) y lo traduce a un '
        'impacto base (1, 3 o 5).'
    )
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = 'CID total'
    hdr[1].text = 'Impacto'
    hdr[2].text = 'Interpretacion'
    rows = [
        ('3 - 4', '1', 'Bajo'),
        ('5 - 7', '3', 'Medio'),
        ('8 - 9', '5', 'Alto'),
    ]
    for r in rows:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = r

    doc.add_paragraph('CAPTURA: Registro de activos y CID (Figura 1)', style='Intense Quote')

    doc.add_heading('3.2 Probabilidad e impacto del escenario', level=2)
    doc.add_paragraph(
        'La probabilidad (1-5) representa la posibilidad de que ocurra un incidente considerando el contexto y '
        'controles existentes. El impacto del escenario se toma del activo (por CID) o se puede ajustar con un '
        'override si el escenario es diferente al impacto general.'
    )

    doc.add_heading('3.3 Formula y clasificacion', level=2)
    doc.add_paragraph('Riesgo (Score) = Probabilidad (1-5) x Impacto (1-5).')
    table2 = doc.add_table(rows=1, cols=3)
    h2 = table2.rows[0].cells
    h2[0].text = 'Score'
    h2[1].text = 'Nivel'
    h2[2].text = 'Accion sugerida'
    for score, lvl, act in [
        ('1 - 5', 'Bajo', 'Monitorear'),
        ('6 - 10', 'Medio', 'Planificar acciones'),
        ('11 - 15', 'Alto', 'Mitigar/transferir con prioridad'),
        ('16 - 25', 'Critico', 'Accion inmediata / escalamiento'),
    ]:
        r = table2.add_row().cells
        r[0].text, r[1].text, r[2].text = score, lvl, act

    doc.add_heading('4. Requisitos implementados (segun consigna)', level=1)
    doc.add_paragraph('A continuacion se muestra como RiskGuard cumple los puntos clave:')

    mapping = [
        ('Identificacion y registro de riesgos',
         'Modulo de Activos, Amenazas, Vulnerabilidades y creacion de Riesgos (escenarios).'),
        ('Valoracion del riesgo',
         'Calculo automatico Score = P x I, usando CID del activo o impacto override. Nivel (Bajo/Medio/Alto/Critico).'),
        ('Tratamiento del riesgo',
         'Estrategia (Mitigar/Transferir/Aceptar/Evitar), controles propuestos, responsable, fecha limite, estado y aprobacion para aceptar.'),
        ('Monitoreo y supervision',
         'Dashboard con KPIs (planes, acciones a tiempo, riesgos reducidos) y registro de incidentes.'),
        ('Comunicacion y consulta',
         'Campo de observaciones/recomendaciones por riesgo y reporte PDF para interesados.'),
        ('Validacion de datos',
         'Formularios con validacion (campos obligatorios, rangos, tipos), y proteccion CSRF en operaciones POST.'),
    ]
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = 'Requisito'
    t.rows[0].cells[1].text = 'Implementacion en RiskGuard'
    for req, impl in mapping:
        row = t.add_row().cells
        row[0].text = req
        row[1].text = impl

    doc.add_paragraph('CAPTURA: Dashboard con KPIs (Figura 2)', style='Intense Quote')
    doc.add_paragraph('CAPTURA: Lista de riesgos y niveles (Figura 3)', style='Intense Quote')
    doc.add_paragraph('CAPTURA: Tratamiento y controles propuestos (Figura 4)', style='Intense Quote')
    doc.add_paragraph('CAPTURA: Registro de riesgo residual (Figura 5)', style='Intense Quote')
    doc.add_paragraph('CAPTURA: Reporte PDF generado (Figura 6)', style='Intense Quote')

    doc.add_heading('5. Arquitectura y tecnologias', level=1)
    doc.add_paragraph('Stack usado:')
    for s in [
        'Backend: Python + Flask (patron MVC ligero).',
        'Base de datos: SQLite (archivo local).',
        'ORM: SQLAlchemy.',
        'Frontend: HTML + Bootstrap 5.',
        'Reportes: ReportLab (PDF).',
    ]:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_paragraph('Arquitectura (descripcion rapida):')
    doc.add_paragraph('Vista (Templates) <-> Controlador (Flask routes) <-> Modelo (SQLAlchemy).', style='List Bullet')

    doc.add_heading('6. Modelo de datos (resumen)', level=1)
    doc.add_paragraph('Principales entidades:')
    for s in [
        'Asset: activo con C/I/D e impacto calculado.',
        'Threat y Vulnerability: catalogos.',
        'RiskScenario: escenario de riesgo con P, impacto override, tratamiento y residual.',
        'Control: controles propuestos (muchos a muchos con RiskScenario).',
        'Incident: incidentes relacionados a un riesgo.',
    ]:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_paragraph('ANEXO: Puedes incluir un diagrama ER (captura de DB Browser o similar).', style='Intense Quote')

    doc.add_heading('7. Consideraciones de seguridad', level=1)
    for s in [
        'Validacion de entrada (campos obligatorios, rangos y tipos).',
        'Proteccion CSRF en formularios y eliminaciones.',
        'Separacion de capas (modelo/controlador/vista) para reducir errores y mejorar mantenimiento.',
        'Recomendacion: en un entorno real, agregar autenticacion/roles y cifrado de datos sensibles.',
    ]:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('8. Dilema etico (ejemplo)', level=1)
    doc.add_paragraph(
        'Dilema: el sistema podria almacenar informacion sobre incidentes o debilidades que, si se filtra, '
        'afectaria la reputacion de la organizacion. El equipo debe balancear transparencia interna (para mejorar) '
        'con confidencialidad (para proteger la informacion).'
    )
    doc.add_paragraph(
        'Decision propuesta: limitar acceso por roles, registrar auditoria, y compartir reportes por niveles '
        'segun necesidad (principio de minimo privilegio).'
    )

    doc.add_heading('9. Impacto social/ambiental/economico (breve)', level=1)
    doc.add_paragraph('Social: reduce probabilidad de filtraciones y protege a usuarios/colaboradores.', style='List Bullet')
    doc.add_paragraph('Ambiental: digitaliza registros (menos papel) y facilita seguimiento sin impresiones.', style='List Bullet')
    doc.add_paragraph('Economico: ayuda a priorizar inversiones en controles donde el riesgo es mas alto.', style='List Bullet')

    doc.add_heading('10. Pruebas (checklist)', level=1)
    for s in [
        'Crear 2 activos con diferente CID y verificar impacto calculado.',
        'Crear amenaza y vulnerabilidad.',
        'Crear riesgo con P=4 y verificar score y nivel.',
        'Definir tratamiento, asignar responsable y fecha.',
        'Registrar residual (P/I mas bajos) y verificar que baje de nivel.',
        'Generar reporte PDF y verificar que incluya riesgos y KPIs.',
    ]:
        doc.add_paragraph(s, style='List Number')

    doc.add_heading('11. Conclusiones y recomendaciones', level=1)
    doc.add_paragraph(
        'RiskGuard cumple la consigna al cubrir el ciclo completo: identificacion, valoracion, tratamiento, '
        'riesgo residual, monitoreo y reporte. Como mejora futura se recomienda: autenticacion, multiusuario, '
        'exportacion a Excel y una matriz 5x5 interactiva.'
    )

    out = OUT_DIR / 'Documento_Tecnico_RiskGuard.docx'
    doc.save(out)
    return out


if __name__ == '__main__':
    m = manual_usuario()
    t = documento_tecnico()
    print('Generated:', m)
    print('Generated:', t)
